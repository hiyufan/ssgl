"""Document parser service for extracting text from various file formats."""

import io
import logging
from pathlib import Path

import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """Extract text content from PDF, Word, TXT, and Markdown files."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}

    def parse(self, file_content: bytes, filename: str) -> str:
        """Parse a file and extract text content.

        Args:
            file_content: Raw file bytes
            filename: Original filename with extension

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        if ext == ".pdf":
            return self._parse_pdf(file_content)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_content)
        elif ext in (".txt", ".md", ".markdown"):
            return self._parse_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())

            if not text_parts:
                raise ValueError("No text content found in PDF")

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            if not text_parts:
                raise ValueError("No text content found in Word document")

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise ValueError(f"Failed to parse Word document: {str(e)}")

    def _parse_text(self, content: bytes) -> str:
        """Extract text from plain text or markdown file."""
        try:
            # Detect encoding
            detected = chardet.detect(content)
            encoding = detected.get("encoding", "utf-8") or "utf-8"

            # Try detected encoding first, fallback to utf-8 and gbk
            for enc in [encoding, "utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    text = content.decode(enc)
                    if text.strip():
                        return text.strip()
                except (UnicodeDecodeError, LookupError):
                    continue

            raise ValueError("Could not decode text file with any supported encoding")
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            raise ValueError(f"Failed to parse text file: {str(e)}")

    def get_file_info(self, filename: str, file_size: int) -> dict:
        """Get file metadata."""
        ext = Path(filename).suffix.lower()
        return {
            "filename": filename,
            "extension": ext,
            "file_size": file_size,
            "file_size_mb": round(file_size / (1024 * 1024), 2),
            "supported": ext in self.SUPPORTED_EXTENSIONS,
        }


# Module-level singleton
document_parser = DocumentParser()
